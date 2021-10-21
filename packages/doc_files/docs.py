import docx
from docx.shared import Pt

from tkinter import messagebox

import os

from ..helpers.manipulators import make_safe_filename, sort_titles, remove_punctuation # type: ignore
from ..helpers.checkers import check_dir # type: ignore

def process_all_docs(list_of_dicts: list, doc_name: str, titles: list, folder_name: str):
    """
    Processes all docx files using the process_doc() function, with a hard limit of 5000 rows for the CSV
    """    
    if len(list_of_dicts) > 5000:
        return messagebox.showerror(title="ERROR", message="There are too many lines in the CSV, this program cannot handle over 5000 lines")
    
    for dictionary in list_of_dicts:
        letter_name = process_doc(dictionary, doc_name, titles, folder_name)
        print(f'Creating letter {letter_name}')

def get_doc_words(doc_name: str) -> set:
    """
    Gets the placeholder words from a doc file and returns them as a set
    """
    doc = docx.Document(doc_name)
    text = ' '.join([para.text for para in doc.paragraphs])
    words = text.split()
    # Get clean (no punctuation) titles/placeholder words, but ignore long underlines for signatures
    titles = [remove_punctuation(word) for word in words if (word[0] == '_' and word[1] != '_')]
    # Use set to remove duplicates
    return set(titles)


def process_doc(replacement_dict: dict, doc_name: str, titles: list, folder_name: str):
    """
    Process and replace words in the document, set styles to calibri and 11pt font,
    finally save file based on the replacement words, adding more columns until
    it finds a filename that isn't already in use.
    """
    sorted_titles = sort_titles(titles)
    doc = docx.Document(doc_name)
    for para in doc.paragraphs:
        for title in sorted_titles:
            text = para.text
            text = text.replace(title, replacement_dict[title])
            para.text = text

    style = doc.styles['Normal']
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    column = 1
    safe_name = make_safe_filename(f'{replacement_dict[titles[0]]} {replacement_dict[titles[column]]}')

    while not check_dir(f'{safe_name}.docx', folder_name):
        column += 1
        safe_name += f' {make_safe_filename(replacement_dict[titles[column]])}'
        if column == len(titles) - 1:
            break

    path = os.path.join(folder_name, f'{safe_name}.docx')
    doc.save(path)
    return path


if __name__ == "__main__":
    pass

else: 
    print("importing docs.py")